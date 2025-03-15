"""文書処理モジュール。

様々な形式の文書を読み込み、テキスト抽出、チャンキング、前処理を行います。
"""
import os
import re
from typing import List, Dict, Any, Optional, Iterator
from datetime import datetime
import pypdf
import docx
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document


def load_document(file_path: str) -> Optional[str]:
    """様々な形式の文書を読み込み、テキストを抽出する。

    Args:
        file_path: 文書ファイルのパス

    Returns:
        抽出されたテキスト。失敗した場合はNone。
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        # PDFファイルの処理
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        
        # Word文書の処理
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        
        # Markdownファイルの処理
        elif file_extension in ['.md', '.markdown']:
            return extract_text_from_markdown(file_path)
        
        # テキストファイルの処理
        elif file_extension in ['.txt', '.text', '.csv', '.json']:
            return extract_text_from_text(file_path)
        
        else:
            print(f"未対応のファイル形式: {file_extension}")
            return None
    
    except Exception as e:
        print(f"文書 {file_path} の読み込み中にエラーが発生しました: {e}")
        return None


def extract_text_from_pdf(file_path: str) -> str:
    """PDFファイルからテキストを抽出する。

    Args:
        file_path: PDFファイルのパス

    Returns:
        抽出されたテキスト
    """
    text = ""
    with open(file_path, 'rb') as f:
        pdf = pypdf.PdfReader(f)
        for page in pdf.pages:
            text += page.extract_text() + "\n\n"
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Word文書からテキストを抽出する。

    Args:
        file_path: Word文書のパス

    Returns:
        抽出されたテキスト
    """
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_text_from_markdown(file_path: str) -> str:
    """Markdownファイルからテキストを抽出する。

    Args:
        file_path: Markdownファイルのパス

    Returns:
        抽出されたテキスト
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    # Markdownをプレーンテキストに変換（HTMLタグは削除）
    html = markdown.markdown(md_text)
    text = re.sub(r'<[^>]+>', '', html)
    return text


def extract_text_from_text(file_path: str) -> str:
    """テキストファイルからテキストを抽出する。

    Args:
        file_path: テキストファイルのパス

    Returns:
        抽出されたテキスト
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def preprocess_text(text: str) -> str:
    """テキストの前処理を行う。

    Args:
        text: 前処理するテキスト

    Returns:
        前処理されたテキスト
    """
    # 余分な空白や改行を削除
    text = re.sub(r'\s+', ' ', text)
    
    # 制御文字を削除
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    
    return text.strip()


def split_text_into_chunks(
    text: str, 
    chunk_size: int = 1000, 
    chunk_overlap: int = 200
) -> List[str]:
    """テキストをチャンクに分割する。

    Args:
        text: 分割するテキスト
        chunk_size: チャンクサイズ（デフォルト: 1000）
        chunk_overlap: チャンクのオーバーラップ量（デフォルト: 200）

    Returns:
        テキストチャンクのリスト
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "、", " ", ""]
    )
    
    chunks = text_splitter.split_text(text)
    return chunks


def process_document(
    file_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:
    """文書を処理し、チャンクと関連メタデータに変換する。

    Args:
        file_path: 文書ファイルのパス
        chunk_size: チャンクサイズ（デフォルト: 1000）
        chunk_overlap: チャンクのオーバーラップ量（デフォルト: 200）

    Returns:
        Document（テキストとメタデータ）のリスト
    """
    # 文書の読み込み
    text = load_document(file_path)
    if text is None:
        return []
    
    # テキストの前処理
    processed_text = preprocess_text(text)
    
    # チャンクに分割
    chunks = split_text_into_chunks(processed_text, chunk_size, chunk_overlap)
    
    # ファイル情報の取得
    file_name = os.path.basename(file_path)
    file_extension = os.path.splitext(file_name)[1]
    file_stat = os.stat(file_path)
    created_time = datetime.fromtimestamp(file_stat.st_ctime)
    modified_time = datetime.fromtimestamp(file_stat.st_mtime)
    
    # Documentオブジェクトを作成
    documents = []
    for i, chunk in enumerate(chunks):
        metadata = {
            "source": file_path,
            "file_name": file_name,
            "file_type": file_extension[1:],  # 先頭の'.'を除去
            "chunk_index": i,
            "total_chunks": len(chunks),
            "created_at": created_time.isoformat(),
            "updated_at": modified_time.isoformat(),
        }
        
        documents.append(Document(page_content=chunk, metadata=metadata))
    
    return documents


def process_directory(
    directory_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:
    """ディレクトリ内の全文書を処理する。

    Args:
        directory_path: 文書ディレクトリのパス
        chunk_size: チャンクサイズ（デフォルト: 1000）
        chunk_overlap: チャンクのオーバーラップ量（デフォルト: 200）

    Returns:
        すべての文書から抽出したDocumentのリスト
    """
    documents = []
    
    # サポートされるファイル拡張子
    supported_extensions = ['.pdf', '.doc', '.docx', '.md', '.markdown', '.txt', '.text']
    
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_extension = os.path.splitext(file)[1].lower()
            
            if file_extension in supported_extensions:
                print(f"文書を処理中: {file_path}")
                docs = process_document(file_path, chunk_size, chunk_overlap)
                documents.extend(docs)
    
    print(f"合計 {len(documents)} チャンクを処理しました。")
    return documents


def format_citation(doc: Document) -> str:
    """文書メタデータから引用情報をフォーマットする。

    Args:
        doc: Documentオブジェクト

    Returns:
        フォーマットされた引用情報
    """
    metadata = doc.metadata
    file_name = metadata.get("file_name", "不明なファイル")
    updated_at = metadata.get("updated_at", "不明な日付")
    
    try:
        # 更新日付を日本語形式に変換
        date_obj = datetime.fromisoformat(updated_at)
        formatted_date = date_obj.strftime("%Y年%m月%d日")
    except (ValueError, TypeError):
        formatted_date = updated_at
    
    return f"[出典: {file_name}, 最終更新: {formatted_date}]"